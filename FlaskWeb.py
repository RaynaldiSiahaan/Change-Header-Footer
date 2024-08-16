import os
import zipfile
import io
from flask import Flask, request, render_template, send_file, redirect, url_for
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['PROCESSED_FOLDER'] = 'processed'

def format_run(run, font_name, font_size, bold):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold

    rPr = run._element
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)

def replace_text_in_paragraphs(paragraphs, old_text, new_text, font_name, font_size, bold):
    for paragraph in paragraphs:
        if old_text in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text
                    format_run(inline[i], font_name, font_size, bold)

def replace_text_in_tables(tables, old_text, new_text, font_name, font_size, bold):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, old_text, new_text, font_name, font_size, bold)

def process_docx(doc, header_replacements, footer_replacements):
    for section in doc.sections:
        for replacement in header_replacements:
            element_type = replacement['element_type']
            old_text = replacement['old_text']
            new_text = replacement['new_text']
            font_name = replacement['font_name']
            font_size = int(replacement['font_size'])
            bold = replacement['bold']

            if element_type == 'Paragraph':
                replace_text_in_paragraphs(section.header.paragraphs, old_text, new_text, font_name, font_size, bold)
            elif element_type == 'Table':
                replace_text_in_tables(section.header.tables, old_text, new_text, font_name, font_size, bold)

        for replacement in footer_replacements:
            element_type = replacement['element_type']
            old_text = replacement['old_text']
            new_text = replacement['new_text']
            font_name = replacement['font_name']
            font_size = int(replacement['font_size'])
            bold = replacement['bold']

            if element_type == 'Paragraph':
                replace_text_in_paragraphs(section.footer.paragraphs, old_text, new_text, font_name, font_size, bold)
            elif element_type == 'Table':
                replace_text_in_tables(section.footer.tables, old_text, new_text, font_name, font_size, bold)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'zip_file' not in request.files:
            return redirect(request.url)

        file = request.files['zip_file']
        if file.filename == '':
            return redirect(request.url)

        if file and file.filename.endswith('.zip'):
            filename = secure_filename(file.filename)
            zip_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(zip_path)

            docx_files = []
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(app.config['UPLOAD_FOLDER'])
                docx_files = [f for f in zip_ref.namelist() if f.endswith('.docx')]

            return render_template('process.html', docx_files=docx_files, zip_filename=filename)

    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    # Collect header fields
    header_old_texts = request.form.getlist('header_old_text')
    header_new_texts = request.form.getlist('header_new_text')
    header_element_types = request.form.getlist('header_element_type')
    header_font_names = request.form.getlist('header_font_name')
    header_font_sizes = request.form.getlist('header_font_size')
    header_bolds = request.form.getlist('header_bold')

    # Collect footer fields
    footer_old_texts = request.form.getlist('footer_old_text')
    footer_new_texts = request.form.getlist('footer_new_text')
    footer_element_types = request.form.getlist('footer_element_type')
    footer_font_names = request.form.getlist('footer_font_name')
    footer_font_sizes = request.form.getlist('footer_font_size')
    footer_bolds = request.form.getlist('footer_bold')

    zip_filename = request.form.get('zip_filename')
    processed_files = []

    with zipfile.ZipFile(os.path.join(app.config['UPLOAD_FOLDER'], zip_filename), 'r') as zip_ref:
        for file in zip_ref.namelist():
            if file.endswith('.docx'):
                doc = Document(io.BytesIO(zip_ref.read(file)))
                
                # Process headers
                for i in range(len(header_old_texts)):
                    if i < len(header_new_texts) and i < len(header_element_types):
                        old_text = header_old_texts[i]
                        new_text = header_new_texts[i]
                        element_type = header_element_types[i]
                        font_name = header_font_names[i] if i < len(header_font_names) else 'Calibri'
                        font_size = int(header_font_sizes[i]) if i < len(header_font_sizes) else 11
                        bold = 'bold' in header_bolds[i] if i < len(header_bolds) else False

                        for section in doc.sections:
                            if element_type == 'Paragraph':
                                replace_text_in_paragraphs(section.header.paragraphs, old_text, new_text, font_name, font_size, bold)
                            elif element_type == 'Table':
                                replace_text_in_tables(section.header.tables, old_text, new_text, font_name, font_size, bold)

                # Process footers
                for i in range(len(footer_old_texts)):
                    if i < len(footer_new_texts) and i < len(footer_element_types):
                        old_text = footer_old_texts[i]
                        new_text = footer_new_texts[i]
                        element_type = footer_element_types[i]
                        font_name = footer_font_names[i] if i < len(footer_font_names) else 'Calibri'
                        font_size = int(footer_font_sizes[i]) if i < len(footer_font_sizes) else 11
                        bold = 'bold' in footer_bolds[i] if i < len(footer_bolds) else False

                        for section in doc.sections:
                            if element_type == 'Paragraph':
                                replace_text_in_paragraphs(section.footer.paragraphs, old_text, new_text, font_name, font_size, bold)
                            elif element_type == 'Table':
                                replace_text_in_tables(section.footer.tables, old_text, new_text, font_name, font_size, bold)

                processed_path = os.path.join(app.config['PROCESSED_FOLDER'], file)
                doc.save(processed_path)
                processed_files.append(processed_path)

    processed_zip = io.BytesIO()
    with zipfile.ZipFile(processed_zip, 'w') as zf:
        for file_path in processed_files:
            zf.write(file_path, os.path.basename(file_path))

    processed_zip.seek(0)

    return send_file(processed_zip, as_attachment=True, download_name='processed_files.zip')


if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['PROCESSED_FOLDER'], exist_ok=True)
    app.run(debug=True)
